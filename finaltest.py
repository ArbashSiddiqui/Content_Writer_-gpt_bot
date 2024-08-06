import os
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from tavily import TavilyClient
from docx import Document

class ArticleProcessor:
    def __init__(self, tavily_api_key, openai_api_key):
        # Set up API keys
        os.environ["TAVILY_API_KEY"] = tavily_api_key
        os.environ["OPENAI_API_KEY"] = openai_api_key
        
        # Initialize the LLM
        self.llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.7)
        
        # Define the prompt templates
        self.summary_prompt = PromptTemplate(
            input_variables=["text"],
            template='''find the key points and updated info in the following text
                        :\n\n{text}\n\nSummary:''' )
        
        self.blog_prompt = PromptTemplate(
            input_variables=["summary", "text"],
            template="Based on the summary and text, generate a new detailed blog with at least 800 words and include bold headings and subheadings:\n\n{summary}\n\nBlog Post:"
        )
        
        # Initialize TavilyClient
        self.client = TavilyClient(api_key=tavily_api_key)

    def get_articles_content(self, query, max_results=5):
        # Execute search query with TavilyClient
        response = self.client.search(query, max_results=max_results, include_raw_content=True)
        
        # Prepare results
        articles = []
        content = ""
        urls = []
        
        for result in response['results']:
            url = result['url']
            article_content = result['raw_content']
            articles.append([url, article_content])
            content += article_content + "\n"
            urls.append(url)
        
        return articles, content, urls

    def create_chain(self, prompt_template):
        return LLMChain(
            llm=self.llm,
            prompt=prompt_template
        )

    def summarize_text(self, text):
        chain = self.create_chain(self.summary_prompt)
        summary = chain.run({"text": text})
        return summary

    def generate_blog(self, summary, text):
        chain = self.create_chain(self.blog_prompt)
        blog_post = chain.run({"summary": summary, "text": text})
        return blog_post

    def save_text_to_word(self, text, filename):
        doc = Document()
        doc.add_heading('Blog Post', 0)
        doc.add_paragraph(text)
        doc.save(filename)

    def process_query(self, query):
        # Get articles content
        _, content, _ = self.get_articles_content(query)
        
        # Summarize the text
        summary = self.summarize_text(content)
        print("\nSummary:")
        print(summary)
        
        # Generate a new blog based on the summary
        blog_post = self.generate_blog(summary, content)
        print("\nGenerated Blog Post:")
        print(blog_post)
        
        # Save to Word document with the query as filename
        filename = f"{query}.docx"
        self.save_text_to_word(blog_post, filename)
        print(f"\nBlog post saved as {filename}")

if __name__ == "__main__":
    # Initialize ArticleProcessor with API keys
    tavily_api_key = "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"
    openai_api_key = ""
    
    processor = ArticleProcessor(tavily_api_key, openai_api_key)
    
    # User input for the query
    query = input("Enter the topic name: ")
    processor.process_query(query)
